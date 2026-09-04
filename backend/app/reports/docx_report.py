"""
PRAMAN AI - Legal Metrology Editable Word (DOCX) Report Generator
Generates editable compliance reports using python-docx.
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime

class DOCXReportGenerator:
    @staticmethod
    def generate(inspection_data: dict, output_filepath: str) -> str:
        """
        Generates an editable DOCX compliance report.
        """
        os.makedirs(os.path.dirname(output_filepath), exist_ok=True)
        doc = Document()

        # Document Header
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_title = p_title.add_run("GOVERNMENT OF INDIA\nMINISTRY OF CONSUMER AFFAIRS, FOOD & PUBLIC DISTRIBUTION\nLEGAL METROLOGY ENFORCEMENT WING")
        run_title.bold = True
        run_title.font.size = Pt(13)
        run_title.font.color.rgb = RGBColor(15, 41, 66)

        p_sub = doc.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_sub = p_sub.add_run("PRAMAN AI — STATUTORY PACKAGING COMPLIANCE REPORT\n")
        run_sub.bold = True
        run_sub.font.size = Pt(11)
        run_sub.font.color.rgb = RGBColor(30, 64, 175)

        # Metadata Table
        table = doc.add_table(rows=4, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        rows = table.rows
        
        rows[0].cells[0].text = f"Inspection ID: {inspection_data.get('inspection_id')}"
        rows[0].cells[1].text = f"Date: {datetime.now().strftime('%d-%b-%Y %H:%M')}"
        
        rows[1].cells[0].text = f"Product Name: {inspection_data.get('product_name')}"
        rows[1].cells[1].text = f"Category: {inspection_data.get('category')}"
        
        rows[2].cells[0].text = f"Inspector: {inspection_data.get('inspector_name')}"
        rows[2].cells[1].text = f"Score: {inspection_data.get('overall_score')}/100"
        
        rows[3].cells[0].text = f"Compliance Status: {inspection_data.get('compliance_status')}"
        rows[3].cells[1].text = f"Governing Rules: Legal Metrology Act 2009 & PCR 2011"

        doc.add_paragraph()

        # Section 1: Findings
        h1 = doc.add_heading("1. Regulatory Summary & Recommendations", level=2)
        doc.add_paragraph(f"Summary: {inspection_data.get('decision_summary', 'Inspection complete.')}")
        doc.add_paragraph(f"Recommended Statutory Action: {inspection_data.get('recommended_action', 'None.')}")

        # Section 2: Mandatory Declarations Table
        doc.add_heading("2. Mandatory Packaging Declarations Audit", level=2)
        t_decl = doc.add_table(rows=1, cols=4)
        hdr_cells = t_decl.rows[0].cells
        hdr_cells[0].text = "Declaration Field"
        hdr_cells[1].text = "Detected Value"
        hdr_cells[2].text = "Status"
        hdr_cells[3].text = "Confidence"

        for d in inspection_data.get("declarations", []):
            r_cells = t_decl.add_row().cells
            r_cells[0].text = str(d.get("field_name", ""))
            r_cells[1].text = str(d.get("detected_value") or "—")
            r_cells[2].text = "FOUND" if d.get("is_found") else "MISSING"
            r_cells[3].text = f"{d.get('confidence_score', 0)}%"

        # Section 3: Violations and Detailed Rule Checks
        doc.add_heading("3. Detailed Legal Metrology Rule Evaluations", level=2)
        for r in inspection_data.get("results", []):
            p = doc.add_paragraph()
            r_run = p.add_run(f"• [{r.get('rule_id')}] {r.get('rule_name')} — Status: {r.get('status')} ({r.get('severity')})\n")
            r_run.bold = True
            p.add_run(f"  Reference: {r.get('source_section')} ({r.get('source_document')})\n")
            p.add_run(f"  Finding: {r.get('explanation')}\n")

        # Save docx
        doc.save(output_filepath)
        return output_filepath
